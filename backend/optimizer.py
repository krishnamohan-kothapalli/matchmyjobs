# optimizer.py
import os, re, base64, logging, tempfile
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)
_client = None
_MODEL  = os.environ.get("CLAUDE_MODEL", "claude-haiku-4-5-20251001")

def _get_client():
    global _client
    if _client is None:
        _client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    return _client

# ── Rewrite prompt ────────────────────────────────────────────────────────────
_REWRITE_PROMPT = """You are editing a resume to improve its ATS keyword score. Make surgical keyword additions only — do not rewrite, reformat, or restructure.

TARGET JOB: {job_title}
MISSING KEYWORDS TO ADD: {missing_skills}
KEYWORDS TO EMPHASIZE: {matched_skills}

ORIGINAL RESUME (edit this — preserve ALL formatting marks):
{resume_text}

JOB DESCRIPTION:
{jd_text}

STRICT RULES:

NEVER CHANGE:
- Candidate's name
- Company names (exact spelling, capitalisation, punctuation)
- Job titles they held
- Locations (city, state, country)
- Dates
- University names
- GPA or grades
- Certification names and dates

WHAT TO CHANGE:
- Weave missing keywords into existing bullet text where they genuinely fit
- Reorder bullets within a role so JD-relevant ones come first
- Add missing keywords to the Skills section only if they cannot fit in bullets
- Strengthen the Professional Summary with 2-3 missing keywords max

VOICE RULES — critical:
- Keep every sentence structure exactly as written
- Keep verb choices — "built" stays "built", never change to "architected" or "developed"  
- Keep sentence length — do not expand short bullets
- BANNED words: spearheaded, orchestrated, leveraged, robust, dynamic, synergistic, streamlined, transformative, holistic, revolutionized, catalyzed
- Do not inflate impact — "helped build" stays "helped build"

FORMAT RULES — critical:
- Every bullet point MUST start with "• " (bullet + space)
- Section headings must be on their own line in ALL CAPS
- Job title on its own line
- Company name, dates on the next line, separated by " | "
- Skills section: each category on its own line as "Category: skill1, skill2, skill3"
- Keep exactly one blank line between sections
- No markdown, no asterisks, no bold markers

OUTPUT: Return ONLY the resume text. Start with the candidate's name on line 1."""

def rewrite_resume(resume_text, jd_text, extraction, current_score):
    missing  = ", ".join(extraction.get("missing_skills",      [])[:15]) or "none"
    matched  = ", ".join(extraction.get("matched_skills",      [])[:10]) or "none"
    title    = extraction.get("job_title", "target role")

    def trunc(t, n):
        return t if len(t) <= n else t[:int(n*.72)] + "\n[...]\n" + t[-int(n*.18):]

    # Pre-process resume: ensure bullets have • prefix so Claude sees the pattern
    lines = []
    for line in resume_text.split('\n'):
        stripped = line.strip()
        # Convert any bullet variant to • so Claude preserves it
        if re.match(r'^[\-\*·▪►▸‣]\s+', stripped):
            line = '• ' + re.sub(r'^[\-\*·▪►▸‣]\s+', '', stripped)
        lines.append(line)
    processed_resume = '\n'.join(lines)

    resp = _get_client().messages.create(
        model=_MODEL, max_tokens=4000,
        messages=[{"role": "user", "content": _REWRITE_PROMPT.format(
            job_title=title,
            missing_skills=missing,
            matched_skills=matched,
            current_score=round(current_score),
            resume_text=trunc(processed_resume, 4000),
            jd_text=trunc(jd_text, 2000),
        )}]
    )
    return resp.content[0].text.strip()


# ── DOCX generation ───────────────────────────────────────────────────────────

_SECTION_WORDS = {
    'experience','work experience','professional experience','employment',
    'work history','career history',
    'education','academic background','qualifications',
    'skills','technical skills','core competencies','competencies',
    'summary','professional summary','profile','objective',
    'career objective','career summary','about me',
    'certifications','certificates','licenses',
    'projects','key projects','selected projects',
    'achievements','accomplishments','awards','honors',
    'publications','research','presentations',
    'volunteer','volunteering','community',
    'languages','interests','hobbies','activities',
    'references','contact','contact information',
}

EXPERIENCE_VERBS = re.compile(
    r'^(developed|designed|built|implemented|led|managed|created|executed|performed|'
    r'conducted|delivered|automated|collaborated|supported|identified|assisted|validated|'
    r'worked|participated|spearheaded|coordinated|maintained|established|improved|'
    r'analyzed|ensured|utilized|applied|drove|oversaw|handled|tested|deployed|'
    r'configured|migrated|integrated|optimized|monitored|reported|resolved|reviewed)\b',
    re.I
)

SKILL_CATEGORY_RE = re.compile(r'^([A-Za-z][^:\|\n]{2,40})[:\|]\s*\S')

def _classify(line):
    """Return line type: name|contact|section|job_header|company_date|skill_category|bullet|blank|body"""
    s = line.strip()
    if not s:
        return 'blank'

    # Explicit bullet character
    if re.match(r'^[•\-\*·▪►▸‣]\s+', s):
        return 'bullet'

    # Section heading — ALL CAPS, short, no year
    clean = s.rstrip(':').lower()
    if clean in _SECTION_WORDS:
        return 'section'
    if s.rstrip(':').isupper() and 3 <= len(s.rstrip(':')) <= 40 and not re.search(r'\b20\d\d\b', s):
        return 'section'

    # Skill category line: "Category: item, item" — must have colon/pipe separator
    if SKILL_CATEGORY_RE.match(s) and len(s) < 150 and not re.search(r'\b(20\d\d|19\d\d)\b', s):
        return 'skill_category'

    # Company | dates or tab-separated date line
    if re.search(r'\b(20\d\d|19\d\d|present|current)\b', s, re.I):
        if '|' in s or '\t' in s or re.search(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', s, re.I):
            return 'company_date'
        return 'job_header'

    # Contact info
    if re.search(r'@\w+\.\w+|\b\d{3}[\s.\-]\d{3}[\s.\-]\d{4}\b|linkedin\.com|github\.com|http', s, re.I):
        return 'contact'

    # Experience bullet without bullet character — detect by action verb + sentence length
    if EXPERIENCE_VERBS.match(s) and len(s) > 40 and s.endswith('.'):
        return 'bullet'

    return 'body'


def _rule(doc, color='2E4A3E'):
    """Thin line under section heading."""
    p   = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'),  '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    p._p.get_or_add_pPr().append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    return p


def _r(p, text, bold=False, size=11.0, color=None, italic=False):
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def generate_docx(resume_text):
    """
    Convert plain-text resume to clean, ATS-safe DOCX.
    - Single column, Calibri, standard margins
    - No tables, text boxes, images, or headers/footers
    - Proper bold on name, section headings, job titles, company names
    - Bullets via Word List Bullet style
    - Skill categories: bold label + normal items
    """
    doc = Document()

    # Page setup
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.75)
        sec.left_margin = sec.right_margin = Inches(1.0)

    # Default style
    n = doc.styles['Normal']
    n.font.name = 'Calibri'
    n.font.size = Pt(11)
    n.paragraph_format.space_before = Pt(0)
    n.paragraph_format.space_after  = Pt(0)
    n.paragraph_format.line_spacing = Pt(13.5)

    lines     = resume_text.split('\n')
    name_done = False
    prev_type = None
    # Collapse consecutive blank lines
    cleaned = []
    prev_blank = False
    for l in lines:
        is_blank = not l.strip()
        if is_blank and prev_blank:
            continue
        cleaned.append(l)
        prev_blank = is_blank
    lines = cleaned

    for raw in lines:
        line  = raw.rstrip()
        ltype = _classify(line)

        # ── Name (very first non-blank line) ─────────────────────────
        if not name_done and ltype != 'blank':
            name_done = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _r(p, line.strip(), bold=True, size=17, color=(15, 60, 45))
            p.paragraph_format.space_after = Pt(3)
            prev_type = 'name'
            continue

        # ── Blank line ────────────────────────────────────────────────
        if ltype == 'blank':
            # Only add spacer if not after section heading rule (already has space)
            if prev_type not in ('section', 'blank'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
            prev_type = 'blank'
            continue

        # ── Contact info ──────────────────────────────────────────────
        if ltype == 'contact':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _r(p, line.strip(), size=10, color=(90, 90, 90))
            p.paragraph_format.space_after = Pt(2)
            prev_type = 'contact'
            continue

        # ── Section heading ───────────────────────────────────────────
        if ltype == 'section':
            # Space before new section
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(5)

            p = doc.add_paragraph()
            _r(p, line.strip().rstrip(':').upper(), bold=True, size=11, color=(15, 60, 45))
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            _rule(doc)
            prev_type = 'section'
            continue

        # ── Job title (line by itself, not a date line) ───────────────
        if ltype == 'job_header':
            p = doc.add_paragraph()
            _r(p, line.strip(), bold=True, size=11)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(1)
            prev_type = 'job_header'
            continue

        # ── Company | dates line ──────────────────────────────────────
        if ltype == 'company_date':
            p = doc.add_paragraph()
            # Split on | or \t to bold the company name
            if '|' in line:
                parts = line.split('|', 1)
                _r(p, parts[0].strip(), bold=True, size=10.5)
                _r(p, '  |  ', bold=False, size=10.5, color=(120,120,120))
                _r(p, parts[1].strip(), italic=True, size=10.5, color=(100,100,100))
            elif '\t' in line:
                parts = line.split('\t', 1)
                _r(p, parts[0].strip(), bold=True, size=10.5)
                _r(p, '   ', size=10.5)
                _r(p, parts[1].strip(), italic=True, size=10.5, color=(100,100,100))
            else:
                _r(p, line.strip(), bold=True, size=10.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            prev_type = 'company_date'
            continue

        # ── Skill category line ───────────────────────────────────────
        if ltype == 'skill_category':
            p = doc.add_paragraph()
            s = line.strip()
            # Try explicit colon or pipe separator first
            m = re.match(r'^([^:\|]+)[:\|]\s*(.+)$', s)
            if m:
                _r(p, m.group(1).strip() + ': ', bold=True, size=10.5)
                _r(p, m.group(2).strip(), size=10.5)
            else:
                # Merged: "CategoryNameskill1, skill2" — split at lowercase→Capital transition
                split = re.search(r'([a-z)])([A-Z])', s)
                if split:
                    cat   = s[:split.start()+1]
                    items = s[split.start()+1:]
                    _r(p, cat.strip() + ': ', bold=True, size=10.5)
                    _r(p, items.strip(), size=10.5)
                else:
                    _r(p, s, size=10.5)
            p.paragraph_format.left_indent  = Inches(0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            prev_type = 'skill_category'
            continue

        # ── Bullet point ──────────────────────────────────────────────
        if ltype == 'bullet':
            text = re.sub(r'^[•\-\*·▪►▸‣]\s+', '', line.strip())
            p = doc.add_paragraph(style='List Bullet')
            _r(p, text, size=10.5)
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            prev_type = 'bullet'
            continue

        # ── Body / fallback ───────────────────────────────────────────
        p = doc.add_paragraph()
        _r(p, line.strip(), size=10.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        prev_type = 'body'

    # Save
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name); path = tmp.name
    with open(path, 'rb') as f: data = f.read()
    os.unlink(path)
    return data


# ── Main optimizer ────────────────────────────────────────────────────────────

def optimize_resume(resume_text, jd_text, extraction, original_score, run_analysis_fn, nlp):
    best_text  = resume_text
    best_score = original_score
    iterations = 0

    while best_score < 80 and iterations < 2:
        iterations += 1
        logger.info("Optimizer pass %d — score %.1f", iterations, best_score)
        try:
            rewritten = rewrite_resume(best_text, jd_text, extraction, best_score)
        except Exception as e:
            logger.error("Rewrite failed: %s", e); break
        try:
            result    = run_analysis_fn(rewritten, jd_text, nlp)
            new_score = result.get("score", best_score)
        except Exception as e:
            logger.error("Scoring failed: %s", e); break
        logger.info("Pass %d: %.1f → %.1f", iterations, best_score, new_score)
        if new_score > best_score:
            best_score = new_score; best_text = rewritten
        else:
            break

    docx_b64 = None
    try:
        docx_b64 = base64.b64encode(generate_docx(best_text)).decode('utf-8')
        logger.info("DOCX generated OK")
    except Exception as e:
        logger.error("DOCX failed: %s", e, exc_info=True)

    return {
        "optimized_text": best_text,
        "original_score": round(original_score, 1),
        "new_score":      round(best_score, 1),
        "improvement":    round(best_score - original_score, 1),
        "iterations":     iterations,
        "docx_b64":       docx_b64,
        "target_reached": best_score >= 80,
    }
